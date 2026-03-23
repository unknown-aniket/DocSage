"""
rag/ingestion.py
----------------
Document ingestion pipeline.
Supports PDF, TXT, DOCX. Extracts text and normalises into
LangChain Document objects for downstream chunking.
"""

from __future__ import annotations
import io
from pathlib import Path
from typing import List

from langchain_core.documents import Document

from utils.logger import get_logger

log = get_logger(__name__)


class DocumentIngester:
    """Load raw files and return a list of LangChain Documents."""

    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx", ".md"}

    def ingest(self, file_path: str | Path) -> List[Document]:
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type '{ext}'. "
                f"Supported: {self.SUPPORTED_EXTENSIONS}"
            )

        log.info("ingesting_document", path=str(path), ext=ext)

        if ext == ".pdf":
            return self._load_pdf(path)
        elif ext == ".txt" or ext == ".md":
            return self._load_text(path)
        elif ext == ".docx":
            return self._load_docx(path)

    # ── Loaders ──────────────────────────────────────────────

    def _load_pdf(self, path: Path) -> List[Document]:
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("Install pypdf: pip install pypdf")

        reader = PdfReader(str(path))
        docs = []
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                docs.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": path.name,
                            "page": page_num,
                            "total_pages": len(reader.pages),
                        },
                    )
                )
        log.info("pdf_loaded", source=path.name, pages=len(docs))
        return docs

    def _load_text(self, path: Path) -> List[Document]:
        text = path.read_text(encoding="utf-8", errors="replace")
        return [
            Document(
                page_content=text,
                metadata={"source": path.name, "page": None},
            )
        ]

    def _load_docx(self, path: Path) -> List[Document]:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")

        docx = DocxDocument(str(path))
        paragraphs = [p.text for p in docx.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)
        return [
            Document(
                page_content=full_text,
                metadata={"source": path.name, "page": None},
            )
        ]

    def ingest_bytes(
        self, content: bytes, filename: str
    ) -> List[Document]:
        """Ingest directly from uploaded bytes (used by the API endpoint)."""
        suffix = Path(filename).suffix.lower()
        tmp = Path("/tmp") / filename
        tmp.write_bytes(content)
        try:
            return self.ingest(tmp)
        finally:
            tmp.unlink(missing_ok=True)
